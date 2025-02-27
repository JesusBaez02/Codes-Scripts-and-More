import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def consolidar_excel(carpeta):
    """Lee y concatena todos los archivos Excel en la carpeta."""
    archivos = [f for f in os.listdir(carpeta) if f.endswith('.xlsx') or f.endswith('.xls')]
    df_total = pd.DataFrame()
    
    for archivo in archivos:
        df = pd.read_excel(os.path.join(carpeta, archivo), engine='openpyxl')
        df_total = pd.concat([df_total, df], ignore_index=True)
    
    return df_total

def formatear_fecha(fecha):
    """Convierte la fecha en formato YYYY-MM-DD, eliminando la hora."""
    if pd.notna(fecha):
        return pd.to_datetime(fecha).strftime('%Y-%m-%d')
    return ""

def generar_reporte_v1(df, output_file):
    """Genera un documento Word con la información consolidada (solo texto)."""
    doc = Document()
    doc.add_heading('Reporte de Actividades', level=1)
    trabajadores = df['PERSONAL'].unique()
    
    for trabajador in trabajadores:
        doc.add_heading(trabajador, level=2)
        df_trabajador = df[df['PERSONAL'] == trabajador]
        
        for _, fila in df_trabajador.iterrows():
            p = doc.add_paragraph()
            p.add_run("Fecha: ").bold = True
            p.add_run(formatear_fecha(fila['FECHA']))
            
            p = doc.add_paragraph()
            p.add_run("Actividad: ").bold = True
            p.add_run(str(fila['ACTIVIDAD REALIZADA']))
            
            p = doc.add_paragraph()
            p.add_run("Proveedores Trabajados:\n").bold = True
            p.add_run(str(fila['PROVEEDORES TRABAJADOS:']))
            
            p = doc.add_paragraph()
            p.add_run("Clientes:\n").bold = True
            p.add_run(str(fila['CLIENTES CON LOS QUE SE TRABAJO:']))
            
            p = doc.add_paragraph()
            p.add_run("Estatus: ").bold = True
            p.add_run(str(fila['STATUS (Documentos en Proceso / Documentos Finalizados):']))
    
    doc.save(output_file)
    print(f"Informe generado: {output_file}")

def generar_grafico(df, output_img):
    """Genera un gráfico de actividades por trabajador."""
    conteo = df['PERSONAL'].value_counts()
    plt.figure(figsize=(10, 5))
    conteo.plot(kind='bar', color='skyblue')
    plt.xlabel("Trabajadores")
    plt.ylabel("Cantidad de Actividades")
    plt.title("Actividades por Trabajador")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(output_img)
    plt.close()

def generar_reporte_v2(df, output_file):
    """Genera un documento Word con la información consolidada y gráfico."""
    doc = Document()
    doc.add_heading('Reporte de Actividades', level=1)
    trabajadores = df['PERSONAL'].unique()
    
    for trabajador in trabajadores:
        doc.add_heading(trabajador, level=2)
        df_trabajador = df[df['PERSONAL'] == trabajador]
        
        for _, fila in df_trabajador.iterrows():
            p = doc.add_paragraph()
            p.add_run("Fecha: ").bold = True
            p.add_run(formatear_fecha(fila['FECHA']))
            
            p = doc.add_paragraph()
            p.add_run("Actividad: ").bold = True
            p.add_run(str(fila['ACTIVIDAD REALIZADA']))
            
            p = doc.add_paragraph()
            p.add_run("Proveedores Trabajados:\n").bold = True
            p.add_run(str(fila['PROVEEDORES TRABAJADOS:']))
            
            p = doc.add_paragraph()
            p.add_run("Clientes:\n").bold = True
            p.add_run(str(fila['CLIENTES CON LOS QUE SE TRABAJO:']))
            
            p = doc.add_paragraph()
            p.add_run("Estatus: ").bold = True
            p.add_run(str(fila['STATUS (Documentos en Proceso / Documentos Finalizados):']))
    
    # Agregar gráfico al final del documento
    img_path = "grafico_actividades.png"
    generar_grafico(df, img_path)
    doc.add_page_break()
    doc.add_heading("Resumen Gráfico", level=2)
    doc.add_picture(img_path, width=Inches(5))
    
    doc.save(output_file)
    print(f"Informe generado: {output_file}")

# Ejecutar
carpeta_excel = "C:/Users/gottec/Documents/Reportes_febrero"  # Carpeta donde están los archivos
df_consolidado = consolidar_excel(carpeta_excel)

# Generar versión 1 (solo texto)
output_docx_v1 = "reporte_actividades_v1.docx"
generar_reporte_v1(df_consolidado, output_docx_v1)

# Generar versión 2 (con gráficos)
# output_docx_v2 = "reporte_actividades_v2.docx"
# generar_reporte_v2(df_consolidado, output_docx_v2)